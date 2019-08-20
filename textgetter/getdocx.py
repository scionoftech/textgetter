import os
import glob
import numpy as np
from docx import Document
from os.path import dirname, abspath
import platform
from skimage.io import MultiImage
from PIL import Image
from pdf2image import convert_from_path
import pytesseract


def img_txt_extract(input_files_path, output_files_path, file_extensions, ocr_path=None, verbose=False):
    """ this will extract text from  images files """
    try:
        if platform.system() == 'Windows':
            pytesseract.pytesseract.tesseract_cmd = ocr_path

        for extension in file_extensions:
            files = glob.glob(input_files_path + "/*." + extension)
            ll = len(files)
            for i, file in enumerate(files, start=1):
                #  __verbose_print(os.path.basename(file) + " file text is being extracted....", verbose)
                file_name = os.path.basename(file)

                img = Image.open(file)
                img_str = pytesseract.image_to_string(np.array(img))

                document = Document()
                document.add_paragraph(img_str)
                document.save(output_files_path + os.sep + file_name[:-4] + '.docx')

                __verbose_print(str(i) + " of " + str(ll) + " file(s) completed", verbose)
    except Exception as e:
        print(e)


def tif_txt_extract(input_files_path, output_files_path, ocr_path=None, verbose=False):
    """ this will extract text from  tif files """
    try:
        if platform.system() == 'Windows':
            pytesseract.pytesseract.tesseract_cmd = ocr_path

        files = set(glob.glob(input_files_path + "/*.TIF") + glob.glob(input_files_path + "/*.tif"))

        ll = len(files)
        for i, file in enumerate(files, start=1):
            #  __verbose_print(os.path.basename(file) + " file text is being extracted....", verbose)
            file_name = os.path.basename(file)

            images = MultiImage(file, plugin='pil')
            img_str = ''
            for img in images:
                img_str += pytesseract.image_to_string(Image.fromarray(img))

            document = Document()
            document.add_paragraph(img_str)
            document.save(output_files_path + os.sep + file_name[:-4] + '.docx')

            __verbose_print(str(i) + " of " + str(ll) + " file(s) completed", verbose)
    except Exception as e:
        print(e)


def pdf_txt_extract(input_files_path, output_files_path, ocr_path=None, verbose=False):
    """ this will extract text from  pdf files """
    try:
        if platform.system() == 'Windows':
            pytesseract.pytesseract.tesseract_cmd = ocr_path

        files = set(glob.glob(input_files_path + "/*.PDF") + glob.glob(input_files_path + "/*.pdf"))

        ll = len(files)
        for i, file in enumerate(files, start=1):
            #  __verbose_print(os.path.basename(file) + " file text is being extracted....", verbose)
            file_name = os.path.basename(file)

            if platform.system() == 'Windows':
                images = convert_from_path(file,
                                           poppler_path=dirname(
                                               abspath(__file__)) + os.sep + 'poppler' + os.sep + 'bin')
            else:
                images = convert_from_path(file)

            img_str = ''
            for img in images:
                img_str += pytesseract.image_to_string(Image.fromarray(np.array(img)))

            document = Document()
            document.add_paragraph(img_str)
            document.save(output_files_path + os.sep + file_name[:-4] + '.docx')

            __verbose_print(str(i) + " of " + str(ll) + " file(s) completed", verbose)
    except Exception as e:
        print(e)


def __verbose_print(log, isshow):
    if isshow:
        print(log)


# if __name__ == '__main__':
#     # img_txt_extract('/home/user/Documents/work_files/innoclique/pdf_work/testing',
#     #                 '/home/user/Documents/work_files/innoclique/pdf_work/testing', ['jpg'],
#     #                 verbose=True)
#     # tif_txt_extract('/home/user/Documents/work_files/innoclique/pdf_work/testing',
#     #                 '/home/user/Documents/work_files/innoclique/pdf_work/testing',
#     #                 verbose=True)
#     # pdf_txt_extract('/home/user/Documents/work_files/innoclique/pdf_work/testing',
#     #                 '/home/user/Documents/work_files/innoclique/pdf_work/testing',
#     #                 verbose=True)
